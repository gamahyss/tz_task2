import os
import random
import string
import tempfile
import shutil
import subprocess
import zipfile
import py7zr
from docx import Document
from openpyxl import Workbook
import xlwt
from fpdf import FPDF
import patoolib

# интерфейс в консоле

def generate_string(min_len=10, max_len=50) -> str:
    '''функция генерирует строку из случайных символов. по умолчанию размер строки - случайное число от 10 до
    50, но можно задать и другие числа в качестве минимального и максимального'''

    length = random.randint(min_len, max_len) # определяем размер строки
    return ''.join(random.choices(string.ascii_letters + string.digits + ' ', k=length)) # генерируем строку заданного размера из случайных символов

def generate_lines() -> list:
    '''функция генерирует несколько строк из случайных символов. количество строк - случайное число от 10
    до 50'''

    num_lines = random.randint(10, 50) # определяем количество строк
    return [generate_string() for _ in range(num_lines)] # генерируем случайные строки с помощью функции generate_string в количестве num_lines

def generate_doc(path: str) -> str:
    '''функция создаёт в заданной директории документ .doc с произвольными данными'''

    lines = generate_lines() # генерируем произвольные данные в виде случайных строк
    doc = Document() # создаём пустой документ
    for line in lines:
        doc.add_paragraph(line) # добавляем данные в документ построчно (новая строка - новый абзац)
    filename = generate_string(5, 10) # генерируем название файла как случайную строку, длина названия - случайное число от 5 до 10
    path_and_name = path + f'/{filename}.doc' # создаём строку в виде полного пути к файлу (с именем)
    doc.save(path_and_name) # сохраняем файл в указанную директорию под сгенерированным именем
    return path_and_name

def generate_docx(path: str) -> str:
    '''функция создаёт в заданной директории документ формата docx с произвольными данными'''

    doc = Document() # создаём документ в котором будут храниться данные
    lines = generate_lines() # генерируем случайные строки которые и будут произвольными данными
    for line in lines:
        doc.add_paragraph(line) # построчно добавляем данные в документ (новая строка - новый абзац)
    filename = generate_string(5, 10) # генерируем название файла как случайную строку размером от 5 до 10 символов
    path_and_name = path + f'/{filename}.docx' # создаём строку в виде полного пути к файлу (с именем)
    doc.save(path_and_name) # сохраняем файл в директории переданной в функцию под именем которое сгенерировано выше
    return path_and_name

def generate_xls(path: str) -> str:
    '''функция создаёт в заданной директории документ формата xls с произвольными данными, записывая их в
    первый столбец'''

    wb = xlwt.Workbook() # создаём пустой документ
    ws = wb.add_sheet('Sheet1') # добавляем лист 1
    lines = generate_lines() # генерируем случайные строки
    for i, line in enumerate(lines):
        ws.write(i, 0, line) # построчно добавляем данные в документ (в первый столбец) на лист 1
    filename = generate_string(5, 10) # генерируем название файла как случайную строку длиной от 5 до 10 символов
    path_and_name = path + f'/{filename}.xls' # создаём строку в виде полного пути к файлу (с именем)
    wb.save(path_and_name) # сохраняем файл в переданной директории под именем сгенерированным ранее
    return path_and_name

def generate_xlsx(path: str) -> str:
    '''функция создаёт документ формата xlsx с произвольными данными, записывая их в первый столбец'''

    wb = Workbook() # создаёт пустой документ
    ws = wb.active # получаем активный лист документа
    lines = generate_lines() # генерируем случайные строки
    for i, line in enumerate(lines, start=1): # записываем данные построчно в первый столбец на активном листе
        ws.cell(row=i, column=1, value=line)
    filename = generate_string(5, 10) # генерируем название файла как случайную строку длиной от 5 до 10 символов
    path_and_name = path + f'/{filename}.xlsx' # создаём строку в виде полного пути к файлу (с именем)
    wb.save(path_and_name) # сохраняем в указанной директории под сгенерированным именем
    return path_and_name

def generate_pdf(path: str) -> str:
    '''функция создаёт документ формата pdf с произвольными данными'''

    pdf = FPDF() # создаём пустой pdf
    pdf.add_page() # добавляем страницу
    pdf.set_font('Arial', size=12) # выбираем шрифт для вводимых данных Arial и размер текста 12
    lines = generate_lines() # генерируем случайные строки
    for line in lines:
        pdf.multi_cell(200, 5, line) # записываем данные в файл построчно, ширина ячейки для записи - 200, высота ячейки - 5
    filename = generate_string(5, 10) # генерируем название файла как случайную строку длиной от 5 до 10 символов
    path_and_name = path + f'/{filename}.pdf' # создаём строку в виде полного пути к файлу (с именем)
    pdf.output(path_and_name) # сохраняем файл в указанной директории под сгенерированным именем
    return path_and_name

def create_zip(filename: str):
    '''функция создаёт зип-архив указанного файла'''

    base_filename = os.path.basename(filename) # извлекаем имя файла без пути
    folder_name = os.path.splitext(base_filename)[0] # имя папки внутри архива - имя файла без расширения
    archive_name = f'{folder_name}.zip' # имя архива точно такое же, как имя папки внутри архива, но с расширением .zip
    with zipfile.ZipFile(archive_name, 'w', zipfile.ZIP_DEFLATED) as zf: # создаём архив со сжатием (с таким же именем, как у папки) в режиме записи
        arcname = os.path.join(folder_name, base_filename) # формируем путь внутри архива (папка/файл)
        zf.write(filename, arcname) # добавляем файл с нужным внутреннем именем

def create_7z(filename: str):
    '''функция создаёт 7z-архив указанного файла'''

    base_filename = os.path.basename(filename) # получаем имя файла без пути
    folder_name = os.path.splitext(base_filename)[0] # имя папки внутри архива - имя файла без расширения
    archive_name = f'{folder_name}.7z' # имя архива точно такое же, как имя папки внутри архива, но с раширением .7z
    with py7zr.SevenZipFile(archive_name, 'w') as archive: # создаём архив 7z в режиме записи
        archive.write(filename, arcname=os.path.join(folder_name, base_filename)) # добавляем в архив путь имя папки/имя файла

def create_rar(filename: str):
    '''функция создаёт rar-архив указанного файла'''

    base_filename = os.path.basename(filename) # находим имя файла без пути
    folder_name = os.path.splitext(base_filename)[0] # имя папки внутри архива - имя файла без расширения
    archive_name = f'{folder_name}.rar' # имя архива - имя папки внутри архива, но с расширением .rar
    with tempfile.TemporaryDirectory() as tmpdir: # создаём временную директорию
        target_subdir = os.path.join(tmpdir, folder_name) # формируем путь к временной подпапке: временная_директория/имя_папки
        os.makedirs(target_subdir) # создаём подпапку
        shutil.copy2(filename, os.path.join(target_subdir, base_filename)) # копируем исходный файл во временную подпапку
        patoolib.create_archive(archive_name, [target_subdir]) # создаём архив из временной подпапки (папка попадёт в архив целиком)

if __name__ == '__main__':
    '''главный блок программы'''

    while True: # пока пользователь не вышел из программы, продолжаем опрашивать его
        actions = ['Сгенерировать хранилище с произвольным содержимым', 'Архивировать файл', 'Выйти'] # список действий, которые может сделать пользователь
        choice = int(input(f'Введите номер действия:\n1. {actions[0]}\n2. {actions[1]}\n3. {actions[2]}')) # выводим действия, пользователь должен ввести номер
        match choice:
            case 1: # если пользователь хочет сгенерировать хранилище с произвольным содержимым
                extensions = {'doc': generate_doc, # словарь соответствия расширений функциям генерации
                              'docx': generate_docx,
                              'xls': generate_xls,
                              'xlsx': generate_xlsx,
                              'pdf': generate_pdf}
                ext = input('Введите расширение генерируемого хранилища (без точки):') # пользователь вводит желаемое расширение
                path = input('Введите директорию, в которой хотите создать файл (полный путь). Если хотите сохранить в папку программы, нажмите d:') # пользователь вводит полный путь до директории
                if path == 'd': # если пользователь выбрал сохранение в папке программы, то находим текущую директорию проекта с помощью os
                    path = os.getcwd()
                full_path = extensions[ext](path) # генерируем файл с тем расширением, которое ввёл пользователь
                print(f'Хранилище с произвольными данными создано. Путь к файлу: {full_path}') # выводим информацию
            case 2: # если пользователь хочет архивировать какой-либо файл
                extensions = {'zip': create_zip, # словарь соответствия расширений архивов функциям архивирования
                              '7z': create_7z,
                              'rar': create_rar}
                path = input('Введите полный путь к файлу, который хотите архивировать:') # пользователь вводит файл, который хочет архивировать
                ext = input('Введите расширение для архива (без точки):') # пользователь вводит желаемое расширение для архива
                extensions[ext](path) # архивируем указанный файл в том расширении которое ввёл пользователь
                print('Архив в создан в директории архивированного файла.') # выводим информацию
            case 3: # если пользователь выбрал выход, прерываем программу
                break
            case _: # если пользователь ввёл что-то другое, продолжаем опрашивать
                continue